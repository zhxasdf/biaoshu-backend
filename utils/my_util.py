import json
import re
import os
import copy
from openai import OpenAI
from html import escape
from zhipuai import ZhipuAI
from docx import Document
from docx.shared import Cm, Pt, RGBColor
from docx.oxml.ns import qn
import markdown
from bs4 import BeautifulSoup
from docx.oxml import OxmlElement

# 写json
def write_json(data, file_path):
    if data != None:
        with open(file_path, 'w', encoding='utf-8') as f:
            json.dump(data, f, ensure_ascii=False, indent=4)
    else:
        print('Json data is None')

# 读json
def open_json(file_path):
    with open(file_path, 'r', encoding='utf-8') as f:
        data = json.load(f)
    return data

# 打开prompt
def open_prompt(path):
    res = open(path, 'r', encoding='utf-8').readlines()
    return ''.join(res)

# 更新prompt
def update_prompt(data_to_be_replaced, raw_prompt):
    new_prompt = copy.deepcopy(raw_prompt)
    for data in data_to_be_replaced:
        new_prompt = new_prompt.replace(data[0], data[1])
    while "\n\n\n" in new_prompt:
        new_prompt = new_prompt.replace("\n\n\n", "\n\n")
    return new_prompt

# 去掉json注释
def strip_gpt4_json_note(result, start_token="```json\n", end_token="```"):
    result = result[result.index(start_token) + len(start_token):]
    result = result[: result.index(end_token)]
    result = result.strip()
    return result

# markdown -> json
def parser_md_to_json(content):
    res_json = None
    from markdown2json import MarkdownToJSON
    parser = MarkdownToJSON(content)
    res_json = parser.markdown_to_ast()
    
    return res_json

# 调用大模型
def use_llm_models(
    input_content='', 
    model_name='glm-4-plus', 
    base_url='',
    api_key='c667f21594404e308c5588d6517ba12a.h3EV8HHYQQsalMYa'
):
    if model_name == 'glm-4-plus' or model_name == 'glm-4-flash-250414':
        client = ZhipuAI(
            api_key=api_key
        )
    else:
        client = OpenAI(
            base_url=base_url,
            api_key=api_key
        )

    response = client.chat.completions.create(
        model=model_name,
        messages=[
            {
                "role": "user",
                "content": input_content
            }
        ],
        stream=True,
    )
    result = ""
    for chunk in response:
        delta = chunk.choices[0].delta.content
        if delta:
            print(delta, end="", flush=True)
            result += delta

    if "```json" in result:
        result = strip_gpt4_json_note(result)
    result = re.sub(r'//.*', '', result)
    
    return result


# 保存markdown文件
def save_markdown_file(markdown_text, file_path):
    if markdown_text != None:
        with open(file_path, "w", encoding="utf-8") as f_md:
            f_md.write(markdown_text)
    else:
        print("Markdown text is None")

# 打开markdown文件
def open_markdown_file(file_path):
    with open(file_path, 'r', encoding='utf-8') as f:
        content = f.read()
    return content

def json_to_markdown(items, level=1):
    md = ""
    for item in items:
        # 保留 section 中的编号
        section_title = item["section"]
        md += f"{'#' * level} {section_title}\n\n"
        for para in item.get("content", []):
            # 如果是图片占位符，单独成行并加特殊标记，便于后续处理
            if isinstance(para, str) and para.strip().startswith("[图片占位："):
                md += f":::image-placeholder\n{para.strip()}\n:::\n\n"
            else:
                md += f"{para}\n\n"
        if "layers" in item and item["layers"]:
            md += json_to_markdown(item["layers"], level + 1)
    return md

def set_numbering_for_heading(paragraph, level):
    """
    为 Heading 段落应用多级编号属性，使用 numId=1。
    模板中需预先定义 numId=1 对应的多级列表样式。
    """
    p = paragraph._p
    pPr = p.get_or_add_pPr()
    numPr = OxmlElement('w:numPr')

    ilvl = OxmlElement('w:ilvl')
    ilvl.set(qn('w:val'), str(level - 1))
    numPr.append(ilvl)

    numId = OxmlElement('w:numId')
    numId.set(qn('w:val'), '1')
    numPr.append(numId)

    pPr.append(numPr)


def md_to_word(md_file,
               docx_file,
               template_path=None):
    """
    将 Markdown 转为 Word 文档：
    - #/##/### 转 Heading 1/2/3 并应用多级编号；
    - #######(>=7级) 将被识别为“逻辑上的更深层级”，优先使用模板中的 Heading 7/8/9，
      若模板未提供则降级为 Heading 6，并根据级别额外增加缩进，但仍设置多级编号的 ilvl；
    - 去除原标题前序号；
    - 段落和列表样式由模板控制；
    - 图片占位符 :::image-placeholder ... ::: 单独成段，若模板中无该样式，则使用 Normal 并设红色。
    模板应预定义：Normal, Heading1-6, List Bullet, 且 numId=1 的多级列表样式。
    """
    # 读取 Markdown
    with open(md_file, 'r', encoding='utf-8') as f:
        md_text = f.read()

    # 替换图片占位符为 div
    md_text = re.sub(
        r':::image-placeholder\n(.*?)\n:::',
        lambda m: f'<div class="image-placeholder">{m.group(1).strip()}</div>',
        md_text,
        flags=re.DOTALL
    )
    # 将 Markdown 中 7 级及以上标题预处理为带 data-level 的 h6，保留实际层级信息
    # 例：####### Title -> <h6 data-level="7">Title</h6>
    def _h7_plus_repl(m):
        hashes = m.group(1)
        level = len(hashes)
        text = m.group(2).strip()
        return f'<h6 data-level="{level}">{escape(text)}</h6>'

    md_text = re.sub(r'^(#{7,})\s*(.+)$', _h7_plus_repl, md_text, flags=re.MULTILINE)

    html = markdown.markdown(md_text)

    # 加载模板或新建文档
    doc = Document(template_path) if template_path else Document()

    # 工具：检查样式是否存在
    def has_style(name):
        return name in [s.name for s in doc.styles]

    # 添加段落并设置样式，支持红色
    def add_paragraph(text, style_name, red=False):
        if has_style(style_name):
            p = doc.add_paragraph(text, style=style_name)
        else:
            p = doc.add_paragraph(text, style='Normal')
            if red:
                run = p.runs[0]
                run.font.color.rgb = RGBColor(255, 0, 0)
        return p

    soup = BeautifulSoup(html, 'html.parser')
    for el in soup.contents:
        tag = el.name or ''
        if re.match(r'h[1-6]', tag):
            lvl = int(tag[1])
            raw = el.get_text()
            # 去除序号前缀
            text = re.sub(r'^\s*\d+(?:[\.\-\)]\d+)*[\.\-\)]?\s*', '', raw).strip()
            # 若有 data-level，说明原始是 7 级及以上（经预处理保存在 h6 上）
            data_level = el.get('data-level')
            real_lvl = int(data_level) if data_level else lvl

            # 选择样式：优先使用模板中对应级别（若存在），否则降级为 Heading 6
            target_style_lvl = real_lvl if data_level and real_lvl <= 9 else lvl
            style_name_exact = f'Heading {target_style_lvl}'
            from docx.shared import Inches

            if has_style(style_name_exact):
                p = doc.add_paragraph(text, style=style_name_exact)
            else:
                # 模板无 Heading 7/8/9 时，退化为 Heading 6
                p = doc.add_paragraph(text, style='Heading 6' if has_style('Heading 6') else 'Normal')
                # 为 >6 级的情况添加缩进，体现层级
                if real_lvl > 6:
                    p.paragraph_format.left_indent = Inches((real_lvl - 6) * 0.3)

            # 始终设置多级编号层级（ilvl 0-8），若模板的 numId=1 支持则会生效
            ilvl = max(0, min(8, real_lvl - 1))
            set_numbering_for_heading(p, ilvl + 1)
        elif tag == 'div' and 'image-placeholder' in el.get('class', []):
            # 若模板中无 image-placeholder，则 red=True
            add_paragraph(el.get_text(), 'image-placeholder', red=True)
        elif tag == 'p':
            add_paragraph(el.get_text(), 'Normal')
        elif tag == 'ul':
            for li in el.find_all('li', recursive=False):
                add_paragraph(li.get_text(), 'List Bullet')

    doc.save(docx_file)