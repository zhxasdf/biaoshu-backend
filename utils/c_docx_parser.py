import os
import re
import pandas as pd
from utils.my_util import *
from docling.document_converter import DocumentConverter

def extract_text(node):
    """递归提取节点及其子节点中的文本内容。"""
    if node["type"] == "text":
        return node.get("value", "")
    text_parts = []
    for child in node.get("children", []):
        text_parts.append(extract_text(child))
    return "".join(text_parts)

def process_list(node):
    """
    将 list 节点转换为列表，每个 list_item 一项，文本前可加编号。
    返回一个字符串列表，例如 ["1. 项目A", "2. 项目B", ...]。
    """
    items = []
    for idx, item in enumerate(node.get("children", []), start=1):
        def extract_from_list_item(n):
            if n["type"] == "text":
                return n.get("value", "")
            parts = []
            for c in n.get("children", []):
                parts.append(extract_from_list_item(c))
            return "".join(parts)
        text = extract_from_list_item(item).strip()
        if text:
            items.append(f"{idx}. {text}")
    return items

def process_table(node):
    """
    将 table 节点解析为字典，包含 headers 列表和 rows 二维列表。
    """
    headers = []
    rows = []
    
    for child in node.get("children", []):
        if child["type"] == "table_head":
            # 处理表头，假设只有一行或多行表头
            for row in child.get("children", []):
                if row["type"] == "table_row":
                    header_cells = []
                    for cell in row.get("children", []):
                        text = extract_text(cell).strip()
                        header_cells.append(text)
                    headers = header_cells
        elif child["type"] == "table_body":
            # 处理表体，可能多行
            for row in child.get("children", []):
                if row["type"] == "table_row":
                    row_cells = []
                    for cell in row.get("children", []):
                        text = extract_text(cell).strip()
                        row_cells.append(text)
                    rows.append(row_cells)
    return {"type": "table", "headers": headers, "rows": rows}

class DocxParser:
    def __init__(self, config):
        self.config = config
        self.docx_file_path = config.FILE_CONFIG.docx_file_path
        self.save_path = config.FILE_CONFIG.save_path
        self.tool_name = config.tool_name
        self.single_spec_data = config.single_spec_data     # 技术规范是单个文件
        self.single_sepc_data_path = config.single_spec_data_path   # 单个技术规范文件路径
        self.single_score_table = config.single_score_table # 评分表是单个文件
        self.single_score_table_path = config.single_score_table_path   # 单个评分表文件路径
        
        self.step1_md_save_path = f'{self.save_path}/step1_docling_res.md'
        self.tables_data_save_path = f'{self.save_path}/step1_tables_data.json'
        self.step2_md2json_save_path = f'{self.save_path}/step2_md2json_res.json'
        self.step3_post_process_res_json_save_path = f'{self.save_path}/step3_post_process_res.json'

        self.progress_file = self.config.progress_file
        self.normal_log_file = self.config.normal_log_file

        self.completed_steps = 0

    def _write_progress(self):
        self.completed_steps += 1
        with open(self.progress_file, "w", encoding="utf-8") as f:
            f.write(f"{self.completed_steps}/{self.config.total_steps}\n")

    # 解析docx -> markdown
    def parser_docx_to_md(self, path, tool_name='docling'):
        md = None
        # 工具解析docx
        if tool_name == 'docling':
            converter = DocumentConverter()
            result = converter.convert(path)

            md = result.document.export_to_markdown()
            if not isinstance(md, str):
                md = str(md)
        else:
            md = None

        return md

    # 提取docx中表格docx -> json
    def extract_tables(self):
        from docx import Document
        import pandas as pd
        document = Document(self.docx_file_path)
        tables_data = []
        for table in document.tables:
            rows = []
            for row in table.rows:
                cells = [cell.text.strip() for cell in row.cells]
                rows.append(cells)
            df = pd.DataFrame(rows)
            df.columns = df.iloc[0]       # 用第一行当表头
            df = df[1:].reset_index(drop=True)
            tables_data.append(df.to_dict(orient='records'))
        return tables_data

    def parse_content_headers(self, content_list):
        """
        将 content 中的以 # 开头的标题解析为层级结构，输出如下三元组：
        (top_sections, parent_content, has_headers)
        - top_sections: 由 content 中的标题构建出的顶层子 section 列表（用于追加到父节点的 layers）
        - parent_content: 父节点最终保留的 content（若解析出标题，则清空；若未解析出标题，保持原样）
        - has_headers: 是否解析到了标题
        规则：
        - 普通文本/表格在遇到最近的标题后，归属于最近的标题的 content；在第一个标题出现前的内容，将在第一个标题创建后合并到该标题 content。
        - 使用 # 的数量表示层级深度，按最小出现层级作为基准层级，构建相对嵌套。
        - 只让叶子节点保留 content：构建完成后，对于有子层级的节点，其 content 只在无法继续细分时才会被填充。
        """
        if not content_list:
            return [], [], False

        header_regex = re.compile(r"^\s*(#{1,})\s*(.+?)\s*$")

        # 临时存放父级在第一个标题出现前的内容
        preface_buffer = []
        # 由 content 中的标题构建出来的顶层 sections
        top_sections = []
        # 当前层级的栈（每个元素: dict(section, content, layers, _level)）
        section_stack = []
        base_level = None
        has_headers = False

        def start_new_section(level: int, title: str):
            nonlocal base_level, section_stack, top_sections, has_headers
            has_headers = True
            if base_level is None:
                base_level = level
            rel_level = max(1, level - base_level + 1)

            # 回退至对应相对层级
            while len(section_stack) >= rel_level:
                section_stack.pop()

            new_section = {"section": title.strip(), "content": [], "layers": [], "_level": rel_level}
            if section_stack:
                section_stack[-1]["layers"].append(new_section)
            else:
                top_sections.append(new_section)
            section_stack.append(new_section)

            # 将第一个标题前的内容并入当前新节的 content
            nonlocal_preface = preface_buffer[:]  # 复制
            if nonlocal_preface:
                section_stack[-1]["content"].extend(nonlocal_preface)
                preface_buffer.clear()

        for item in content_list:
            if isinstance(item, str):
                m = header_regex.match(item)
                if m:
                    level = len(m.group(1))
                    title = m.group(2)
                    start_new_section(level, title)
                else:
                    if section_stack:
                        section_stack[-1]["content"].append(item)
                    else:
                        preface_buffer.append(item)
            elif isinstance(item, dict) and ("section" in item):
                # 遇到已经是 section 形态的字典，直接作为层级节点提升到 layers
                has_headers = True
                # 首个标题型节点出现时，合并前言到它的 content
                target_parent_layers = None
                if section_stack:
                    target_parent_layers = section_stack[-1]["layers"]
                else:
                    target_parent_layers = top_sections

                # 复制以避免引用共享副作用
                new_section = {
                    "section": item.get("section", "").strip(),
                    "content": list(item.get("content", [])),
                    "layers": list(item.get("layers", []))
                }
                # 将 preface 并入此节的 content
                if preface_buffer:
                    new_section["content"] = list(preface_buffer) + new_section["content"]
                    preface_buffer.clear()

                target_parent_layers.append(new_section)
                # 将其视作当前层级的最新节点
                # 由于未知其层级深度，按与当前栈同级追加（不改变 base_level），仅将其作为当前父层的最后一个子节点
                # 不入栈，避免破坏相对层级结构；其内部层级会在递归阶段清理
            else:
                # 表格等结构化对象
                if section_stack:
                    section_stack[-1]["content"].append(item)
                else:
                    preface_buffer.append(item)

        # 清理辅助字段，并确保“只有叶子节点有 content”
        def cleanup_and_enforce_leaf_content(node_dict):
            for child in node_dict.get("layers", []):
                cleanup_and_enforce_leaf_content(child)
            # 若有子层级，不保留自身 content
            if node_dict.get("layers"):
                node_dict["content"] = []
            # 删除内部辅助字段
            node_dict.pop("_level", None)

        for s in top_sections:
            cleanup_and_enforce_leaf_content(s)

        # 若解析出了标题，则父节点不保留 content；否则保留原 content
        parent_content = [] if has_headers else content_list

        return top_sections, parent_content, has_headers

    def process_content_headers_recursive(self, node):
        """
        递归处理节点：将 content 中的 # 标题解析为层级并移入当前节点的 layers。
        """
        if 'content' in node and node['content']:
            new_layers, parent_content, has_headers = self.parse_content_headers(node['content'])
            if has_headers:
                # 追加由 content 中提取出来的层级
                node.setdefault('layers', [])
                node['layers'].extend(new_layers)
                # 确保父节点在产生子层级后不再保留 content
                node['content'] = []
            else:
                node['content'] = parent_content

        # 递归处理子层级（包括刚刚追加的）
        if 'layers' in node and node['layers']:
            for child in node['layers']:
                self.process_content_headers_recursive(child)
        
        return node

    # 后处理json -> json
    def post_process(self, nodes, divide_chapter=False):
        """
            将平铺的 AST 节点列表转换为层级 JSON 结构：
            - heading 节点根据 depth 嵌套为 layers
            - paragraph 节点作为单独条目追加到对应节的 content 列表
            - list 节点每一项作为单独条目追加到 content 列表
            - table 节点解析后追加 dict 到 content 列表
            返回最外层的列表，即深度为最小 heading 组成的列表。
        """
        root = {"depth": 0, "layers": []}
        stack = [root]

        for node in nodes:
            node_type = node.get("type")
            if node_type == "heading":
                depth = node.get("depth", 0)
                heading_text = extract_text(node).strip()
                # pop 直到找到更小 depth 的 parent
                while stack and stack[-1]["depth"] >= depth:
                    stack.pop()
                section = {"section": heading_text, "content": [], "layers": [], "depth": depth}
                stack[-1]["layers"].append(section)
                stack.append(section)
            elif node_type == "paragraph":
                if len(stack) > 1:
                    text = extract_text(node).strip()
                    if text:
                        stack[-1]["content"].append(text)
            elif node_type == "list":
                if len(stack) > 1:
                    items = process_list(node)
                    for item in items:
                        stack[-1]["content"].append(item)
            elif node_type == "table":
                if len(stack) > 1:
                    table_dict = process_table(node)
                    stack[-1]["content"].append(table_dict)
            else:
                # 其他类型可根据需要扩展
                pass

        # 清理临时 depth 字段
        def clean(node):
            node.pop("depth", None)
            for child in node.get("layers", []):
                clean(child)

        for top in root["layers"]:
            clean(top)

        final_res = []
        chapter_pattern = r'^(?:\d+\s*)?第[一二三四五六七八九十百]+章\s+(.+)'
        temp_layer = []
        for item in root['layers']:
            if re.match(chapter_pattern, item['section']):
                if temp_layer != []:
                    final_res[-1]['layers'].extend(temp_layer)
                final_res.append(item)
                temp_layer = []
            else:
                temp_layer.append(item)

        if divide_chapter:
            if temp_layer != []:
                final_res[-1]['layers'].extend(temp_layer)
        else:
            final_res = {
                'section': '技术规范书',
                'content': [],
                'layers': []
            }
            final_res['layers'].extend(temp_layer)

        # 后处理：将 content 中的标题解析为层级，并移动到 layers
        if isinstance(final_res, list):
            for item in final_res:
                self.process_content_headers_recursive(item)
        else:
            self.process_content_headers_recursive(final_res)
        
        return final_res

    # 提取单个技术规范
    def extract_single_spec_data(self, input_path):
        raw_md = self.parser_docx_to_md(path=input_path, tool_name=self.tool_name)
        raw_json = parser_md_to_json(raw_md)
        post_process_json = self.post_process(raw_json['children'], divide_chapter=False)
        return post_process_json

    # 提取单个评分表
    def extract_single_score_table(self, input_path):
        if input_path.endswith('.xlsx'):
            df = pd.read_excel(input_path, sheet_name=0)
            json_data = df.to_dict(orient='records')
            return json_data
        elif input_path.endswith('.docx'):
            from docx import Document
            document = Document(input_path)
            tables_data = []
            for table in document.tables:
                rows = []
                for row in table.rows:
                    cells = [cell.text.strip() for cell in row.cells]
                    rows.append(cells)
                df = pd.DataFrame(rows)
                df.columns = df.iloc[0]       # 用第一行当表头
                df = df[1:].reset_index(drop=True)
                tables_data.append(df.to_dict(orient='records'))
            return tables_data
        else:
            print('输入文件格式错误')
            return None

    # 解析docx主流程
    def debug_main_process(self):
        print('开始处理文档')

        print('step1: docx转markdown...')
        if os.path.exists(self.step1_md_save_path):
            step1_md = open_markdown_file(self.step1_md_save_path)
        else:
            step1_md = self.parser_docx_to_md(path=self.docx_file_path, tool_name=self.tool_name) if self.docx_file_path != '' else None
            save_markdown_file(step1_md, self.step1_md_save_path)

        print('  step1.1: 提取表格数据...')
        if os.path.exists(self.tables_data_save_path):
            step1_tables_data = open_json(self.tables_data_save_path)
        else:
            step1_tables_data = self.extract_tables() if self.docx_file_path != '' else None
            write_json(step1_tables_data, self.tables_data_save_path)

        print('step2: markdown转json...')
        if os.path.exists(self.step2_md2json_save_path):
            step2_json = open_json(self.step2_md2json_save_path)
        else:
            step2_json = parser_md_to_json(step1_md) if step1_tables_data != None else None
            write_json(step2_json, self.step2_md2json_save_path)

        print('step3: 规则后处理...')
        if os.path.exists(self.step3_post_process_res_json_save_path):
            step3_json = open_json(self.step3_post_process_res_json_save_path)
        else:
            step3_json = self.post_process(step2_json['children']) if step2_json != None else None
            write_json(step3_json, self.step3_post_process_res_json_save_path)

        final_res = {
            'tables_data': step1_tables_data,   # 主招标文件中所有的表格汇总
            'post_process_data': step3_json     # 规则后处理后的主招标文件的json
        }

        # 如果技术规范是单个文件,直接提取
        if self.single_spec_data:
            spec_data = self.extract_single_spec_data(self.single_sepc_data_path)
            final_res.update({'spec_data': spec_data})  # 提取技规并后处理的json

        if self.single_score_table:
            score_table = self.extract_single_score_table(self.single_score_table_path)
            final_res.update({'score_table': score_table})  # 评分表的json

        return final_res

    # 写日志
    def write_normal_log(self, message):
        with open(self.normal_log_file, "a+", encoding="utf-8") as f:
            f.write(message)
            f.write("\n")

    # 解析docx主流程
    def main_process(self):
        print('  step1: docx转markdown...')
        self.write_normal_log(
            '  step1: docx转markdown...'
        )
        step1_md = self.parser_docx_to_md(path=self.docx_file_path, tool_name=self.tool_name) if self.docx_file_path != '' else None

        print('    step1.1: 提取表格数据...')
        self.write_normal_log(
            '    step1.1: 提取表格数据...'
        )
        step1_tables_data = self.extract_tables() if self.docx_file_path != '' else None

        print('  step2: markdown转json...')
        self.write_normal_log(
            '  step2: markdown转json...'
        )
        step2_json = parser_md_to_json(step1_md) if step1_tables_data != None else None

        print('  step3: 规则后处理...')
        self.write_normal_log(
            '  step3: 规则后处理...'
        )
        step3_json = self.post_process(step2_json['children']) if step2_json != None else None

        final_res = {
            'tables_data': step1_tables_data,   # 主招标文件中所有的表格汇总
            'post_process_data': step3_json     # 规则后处理后的主招标文件的json
        }

        # 如果技术规范是单个文件,直接提取
        if self.single_spec_data:
            self.write_normal_log(
                '  step3.1: 解析单个技术规范文件...'
            )
            spec_data = self.extract_single_spec_data(self.single_sepc_data_path)
            final_res.update({'spec_data': spec_data})  # 提取技规并后处理的json

        if self.single_score_table:
            self.write_normal_log(
                '  step3.1: 解析单个评分表...'
            )
            score_table = self.extract_single_score_table(self.single_score_table_path)
            final_res.update({'score_table': score_table})  # 评分表的json

        self._write_progress()
        return final_res













